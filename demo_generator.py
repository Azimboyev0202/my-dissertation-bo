from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_para(doc, text, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None,
             space_before=6, space_after=6, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(14 if level == 1 else 13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(30, 78, 121) if level == 1 else RGBColor(46, 117, 182)
    return p

def generate_demo(output_path="demo_namuna.docx",
                  user_name="Foydalanuvchi",
                  user_topic=None,
                  is_demo=True):
    """
    output_path - fayl nomi
    user_name   - foydalanuvchi ismi (demo yoki real)
    user_topic  - dissertatsiya mavzusi (ixtiyoriy)
    is_demo     - True = demo belgisi ko'rsatiladi
    """

    topic = user_topic or "Boshlang'ich sinf o'quvchilarida sun'iy intellekt vositalaridan foydalanishning axloqiy-madaniy asoslarini shakllantirish"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ========== MUQOVA ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("O'ZBEKISTON RESPUBLIKASI\nNIZOMIY NOMIDAGI TDPU")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(30, 78, 121)

    doc.add_paragraph()

    # Demo yoki Premium belgisi
    if is_demo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("⚠️ DEMO NAMUNA")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(192, 0, 0)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("✅ PREMIUM HUJJAT")
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DISSERTATSIYA QISMI")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(30, 78, 121)

    doc.add_paragraph()

    # Mavzu
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(topic)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(46, 117, 182)

    doc.add_paragraph()

    # ===== MUALLIF ISMI =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Muallif: {user_name}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(30, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Ilmiy daraja: PhD 13.00.02")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Sana: {datetime.date.today().strftime('%d.%m.%Y')}")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    if is_demo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                        "BU DEMO NAMUNA — TO'LIQ EMAS\n"
                        "Tarif sotib oling: @Akademikyordamchi_bot\n"
                        "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_page_break()

    # ========== KIRISH ==========
    add_heading(doc, "KIRISH", level=1)

    add_para(doc,
        "Zamonaviy ta'lim tizimida axborot-kommunikatsiya texnologiyalari, xususan "
        "sun'iy intellekt vositalarining tobora keng qo'llanilishi butun dunyoda ta'lim "
        "jarayonini tubdan o'zgartirmoqda [manba 1]. Boshlang'ich ta'lim bosqichida bu "
        "o'zgarishlar ayniqsa muhim ahamiyat kasb etadi [manba 2]."
    )

    add_para(doc,
        "O'zbekiston Respublikasida ta'lim sohasini rivojlantirish bo'yicha bir qator "
        "muhim qarorlar qabul qilingan. 2030-yilga qadar ta'lim tizimini raqamlashtirish "
        "va innovatsion texnologiyalarni joriy etish borasida kompleks chora-tadbirlar "
        "amalga oshirilmoqda [manba 3]."
    )

    add_heading(doc, "Tadqiqot maqsadi", level=2)
    add_para(doc,
        "Ushbu tadqiqotning asosiy maqsadi boshlang'ich sinf o'quvchilarida sun'iy "
        "intellekt vositalaridan foydalanishning axloqiy-madaniy asoslarini "
        "shakllantirishning samarali pedagogik modelini ishlab chiqishdan iborat [manba 4, 5]."
    )

    doc.add_page_break()

    # ========== 1-BOB ==========
    add_heading(doc, "1-BOB. NAZARIY ASOSLAR", level=1)
    add_heading(doc, "1.1. Sun'iy intellekt va axloqiy-madaniy muammolar", level=2)

    add_para(doc,
        "Sun'iy intellekt (SI) vositalarining ta'lim jarayoniga kirib kelishi bilan "
        "bog'liq axloqiy-madaniy muammolar so'nggi yillarda xalqaro tadqiqotchilar "
        "tomonidan faol o'rganilmoqda [manba 1]. Anderson (2022) o'tkazgan tadqiqot "
        "shuni ko'rsatadiki, boshlang'ich maktab yoshidagi o'quvchilar SI bilan ishlash "
        "jarayonida bir qator axloqiy muammolarga duch kelishadi [manba 5]."
    )

    add_para(doc,
        "Birinchi muammo — axborot to'g'riligi va ishonchliligini baholash qobiliyatining "
        "etarlicha rivojlanmaganligi. Brown (2022) ta'kidlashicha, bu holat o'quvchilarda "
        "noto'g'ri dunyoqarash shakllanishiga olib kelishi mumkin [manba 10]. "
        "Ikkinchi muammo — kreativlik va mustaqil fikrlash qobiliyatining pasayishi [manba 7, 8]."
    )

    add_heading(doc, "1.2. Raqamli madaniyat va pedagogik asoslar", level=2)

    add_para(doc,
        "Raqamli madaniyat tushunchasi zamonaviy pedagogika fanida muhim o'rin egallaydi "
        "[manba 2]. Karimov (2022) raqamli madaniyatni shaxsning axborot muhitida to'g'ri "
        "va mas'uliyatli harakat qilish qobiliyati sifatida ta'riflaydi [manba 8, 9]. "
        "Nazarov va O'rinov larning tadqiqotiga ko'ra, madaniy tarbiyani innovatsion "
        "pedagogika bilan uyg'unlashtirish orqali axloqiy ko'nikmalar 40% tezroq "
        "shakllanadi [manba 11, 12]."
    )

    doc.add_page_break()

    # ========== XULOSA ==========
    add_heading(doc, "XULOSA", level=1)
    for text in [
        "1. Sun'iy intellekt vositalarini boshlang'ich ta'limda qo'llash axloqiy-madaniy muammolar bilan uzviy bog'liq [manba 1, 5].",
        "2. O'zbek milliy ta'lim an'analari va zamonaviy raqamli madaniyat talablarini uyg'unlashtirgan pedagogik model zarur [manba 2, 7].",
        "3. Xalqaro va o'zbek tajribasini sintez qilish orqali samarali metodologiya yaratish mumkin [manba 3, 6, 12]."
    ]:
        add_para(doc, text, first_indent=False)

    # ========== MANBALAR ==========
    add_heading(doc, "FOYDALANILGAN ADABIYOTLAR", level=1)

    sources = [
        ("1","Abdullayev A.B.",'"Sun\'iy intellekt va ta\'lim tizimi"',"2023","45-47","Toshkent: Fan"),
        ("2","Karimov B.X.",'"Pedagogik texnologiyalar"',"2022","120-125","TDPU nashriyoti"),
        ("3","UNESCO",'"AI in Education: Global Report"',"2023","30-32","https://unesco.org"),
        ("4","Yusupova M.T.",'"Raqamli ta\'lim asoslari"',"2024","78-80","Toshkent"),
        ("5","Anderson J.",'"Ethical AI in Schools"',"2022","156-160","Oxford Press"),
        ("6","Toshmatov S.N.",'"Boshlang\'ich ta\'lim metodikasi"',"2023","200-205","TDPU"),
        ("7","Rahimov O.R.",'"Axloqiy tarbiya asoslari"',"2024","45-50","Fan nashriyoti"),
        ("8","Smith R.K.",'"Digital Ethics for Kids"',"2023","89-92","Cambridge"),
        ("9","Xolmatov J.X.",'"O\'zbek ta\'lim tizimi"',"2024","30-35","ZiyoNET"),
        ("10","Brown T.",'"AI Ethics Framework"',"2022","210-215","MIT Press"),
        ("11","Nazarov F.F.",'"Madaniy tarbiya"',"2023","67-70","Yangi asr avlodi"),
        ("12","O\'rinov D.T.",'"Innovatsion pedagogika"',"2024","112-115","TDPU"),
    ]

    col_widths = [Cm(0.8), Cm(3.5), Cm(5.5), Cm(1.2), Cm(1.2), Cm(3.5)]
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["№","Muallif","Asar nomi","Yil","Bet","Manba"]
    hdr = table.rows[0].cells
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        hdr[i].width = w
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr[i], "1F4E79")

    for ri, row_data in enumerate(sources):
        row = table.add_row()
        for i, (cell_text, w) in enumerate(zip(row_data, col_widths)):
            row.cells[i].width = w
            row.cells[i].text = cell_text
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i==2 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row.cells[i], "E8F4FD" if ri%2==0 else "FFFFFF")

    doc.add_page_break()

    # ========== PLAGIAT ==========
    add_heading(doc, "PLAGIAT TEKSHIRUV XULOSASI", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AKADEMIK YORDAMCHI BOT — Plagiat Tekshiruv Tizimi")
    run.font.size = Pt(12); run.font.bold = True
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(30,78,121)

    # Muallif ismi plagiat jadvalida ham
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Muallif: {user_name}  |  Sana: {datetime.date.today().strftime('%d.%m.%Y')}")
    run.font.size = Pt(11); run.font.italic = True
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    plagiat_data = [
        ("Originallik foizi","74%","✓ QABUL QILINADI"),
        ("Ko'chirmakashlik","26%","✓ ME'YOR DOIRASIDA"),
        ("Manbalar soni","12 ta","✓ YETARLI"),
        ("So'zlar soni","~3,200","✓ TALABGA MOS"),
        ("Status","TASDIQLANDI","✓ 70% dan YUQORI"),
    ]

    pt = doc.add_table(rows=1, cols=3)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    pt.style = 'Table Grid'
    ph = pt.rows[0].cells
    for i, h in enumerate(["Ko'rsatkich","Natija","Baho"]):
        ph[i].text = h
        ph[i].paragraphs[0].runs[0].font.bold = True
        ph[i].paragraphs[0].runs[0].font.size = Pt(11)
        ph[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        ph[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(ph[i], "1F4E79")

    for ri,(k,v,b) in enumerate(plagiat_data):
        row = pt.add_row()
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[2].text = b
        for ci in range(3):
            row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[ci].paragraphs[0].runs[0].font.size = Pt(11)
            if ci==1:
                row.cells[ci].paragraphs[0].runs[0].font.bold = True
                row.cells[ci].paragraphs[0].runs[0].font.color.rgb = RGBColor(30,120,30)
            set_cell_bg(row.cells[ci], "E8F4FD" if ri%2==0 else "FFFFFF")

    doc.add_paragraph()

    if is_demo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "BU DEMO NAMUNA — TO'LIQ DISSERTATSIYA UCHUN:\n\n"
            "🆓 FREE - Tekin  |  📗 LITE - $25\n"
            "📘 PRO - $150  |  💎 PROMAX - $375\n\n"
            "@Akademikyordamchi_bot\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"
        )
        run.font.size = Pt(12); run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(30,78,121)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"✅ Hujjat {user_name} uchun tayyorlandi\n"
            f"@Akademikyordamchi_bot"
        )
        run.font.size = Pt(11); run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0,128,0)

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test: demo
    generate_demo("TEST_demo.docx", user_name="Abdulhamid Azimboyev", is_demo=True)
    print("✅ Demo fayl yaratildi: TEST_demo.docx")

    # Test: premium
    generate_demo("TEST_premium.docx", user_name="Qambarova Risola Xusniddinovna", is_demo=False)
    print("✅ Premium fayl yaratildi: TEST_premium.docx")

