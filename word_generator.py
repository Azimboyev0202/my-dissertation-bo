# word_generator.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List

class WordGenerator:
    """Dissertatsiya bandini Word formatda yaratish"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Asosiy uslublarni o'rnatish"""
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
    
    def add_header(self, band_number: str, title: str, author: str = ""):
        """Header qo'shish"""
        
        # Sarlavha
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"BAND {band_number}")
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Nomi
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Muallif
        if author:
            author_para = self.doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(f"Muallif: {author}")
            run.font.size = Pt(11)
            run.font.italic = True
        
        self.doc.add_paragraph()  # Spacing
    
    def add_content(self, content: str):
        """Asosiy matnni qo'shish"""
        para = self.doc.add_paragraph(content)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Inches(0.5)
    
    def add_bot_note(self, note: str, note_type: str = "info"):
        """Bot izohlarini qo'shish"""
        para = self.doc.add_paragraph()
        
        # Ramka
        shading_elm = para._element.get_or_add_pPr()
        shading = shading_elm.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shading is None:
            from docx.oxml import parse_xml
            shading_xml = '<w:shd {} w:fill="FFFACD"/>'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            )
            shading_elm.append(parse_xml(shading_xml))
        
        # Matn
        icon = "ℹ️" if note_type == "info" else "⚠️"
        run = para.add_run(f"{icon} [BOT IZOHI] {note}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(184, 134, 11)
        
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.right_indent = Inches(0.25)
    
    def add_sources(self, sources: List[Dict]):
        """Manbalar ro'yxati"""
        self.doc.add_paragraph()
        
        heading = self.doc.add_paragraph()
        run = heading.add_run("MANBALAR")
        run.font.bold = True
        run.font.size = Pt(11)
        
        for source in sources:
            source_text = f"[{source['number']}] {source['author']}. {source['title']}. {source['year']}"
            para = self.doc.add_paragraph(source_text, style='List Number')
    
    def add_recommendations(self, recommendations: List[str]):
        """Tavsiyalar"""
        self.doc.add_paragraph()
        
        heading = self.doc.add_paragraph()
        run = heading.add_run("RAHBAR UCHUN TAVSIYALAR")
        run.font.bold = True
        run.font.size = Pt(11)
        
        for i, rec in enumerate(recommendations, 1):
            para = self.doc.add_paragraph(rec, style='List Number')
    
    def save(self, filename: str):
        """Word faylni saqlash"""
        self.doc.save(filename)
        return filename
    
    def get_document(self):
        """Document objekt olish"""
        return self.doc


def create_band_document(band_number: str, title: str, content: str,
                        plagiat_analysis: Dict, author: str = "") -> Document:
    """Band uchun Word dokument yaratish"""
    
    generator = WordGenerator()
    
    # Header
    generator.add_header(band_number, title, author)
    
    # Matn
    generator.add_content(content)
    
    # Bot tavsiyalari
    if plagiat_analysis.get('recommendations'):
        generator.doc.add_paragraph()
        for rec in plagiat_analysis['recommendations']:
            generator.add_bot_note(rec, "info")
    
    # Plagiat hisobi
    generator.doc.add_paragraph()
    stats_para = generator.doc.add_paragraph()
    run = stats_para.add_run(
        f"Originallik: {plagiat_analysis['originallikFoizi']}% | "
        f"Status: {plagiat_analysis['statusBadge']}"
    )
    run.font.size = Pt(10)
    run.font.bold = True
    
    # Vaqt
    generator.doc.add_paragraph()
    footer = generator.doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Tekshiruv: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    return generator.doc
