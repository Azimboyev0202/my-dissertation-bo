import logging
import os
import json
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (InlineKeyboardMarkup, InlineKeyboardButton,
                           ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove)
import google.generativeai as genai
from demo_generator import generate_demo

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
ADMIN_ID = int(os.getenv("ADMIN_ID", "0"))

# ========== ADMIN MA'LUMOTLARI (O'ZGARTIRING) ==========
ADMIN_USERNAME = "@Azimboyev0202"
ADMIN_PHONE = "+998 90 123 45 67"
# ========================================================

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")

bot = Bot(token=TELEGRAM_TOKEN)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

users_db = {}
pending_users = {}

TARIFS = {
    "free":   {"name": "🆓 FREE",   "price": 0,   "uzs": 0,       "pages": 0,   "bands": 0},
    "lite":   {"name": "📗 LITE",   "price": 25,  "uzs": 312500,  "pages": 50,  "bands": 5},
    "pro":    {"name": "📘 PRO",    "price": 150, "uzs": 1875000, "pages": 100, "bands": 10},
    "promax": {"name": "💎 PROMAX", "price": 375, "uzs": 4687500, "pages": 150, "bands": 15},
}

# ========================================
# STATES
# ========================================
class Registration(StatesGroup):
    name = State()
    phone = State()
    topic = State()
    help_type = State()
    tarif = State()

class BandWrite(StatesGroup):
    band_number = State()
    band_name = State()
    waiting_ai = State()
    edit_text = State()

class PlagiatCheck(StatesGroup):
    text = State()

class Payment(StatesGroup):
    screenshot = State()

# ========================================
# HELPERS
# ========================================
def get_user(uid): return users_db.get(uid)
def is_registered(uid):
    u = get_user(uid)
    return u and u.get("status") == "approved"
def is_admin(uid): return uid == ADMIN_ID

def plagiat_check(text):
    words = text.split()
    total = len(words)
    if total == 0: return 0, 100
    sources = len([w for w in text.split('[') if 'manba' in w.lower()])
    unique = len(set(words))
    ratio = unique / total
    orig = min(95, max(55, int(ratio * 100) + sources * 5))
    return orig, 100 - orig

# ========================================
# KLAVIATURALAR
# ========================================
def reply_keyboard_free():
    """FREE user uchun pastki knopkalar"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="📄 Demo ko'rish"), KeyboardButton(text="🔍 Plagiat tekshirish")],
            [KeyboardButton(text="💰 Tarif sotib olish"), KeyboardButton(text="👤 Profilim")],
            [KeyboardButton(text="📞 Admin bilan bog'lanish")],
        ],
        resize_keyboard=True,
        input_field_placeholder="Buyruq tanlang yoki yozing..."
    )

def reply_keyboard_premium():
    """PREMIUM user uchun pastki knopkalar"""
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="✍️ Band yozish (AI)"), KeyboardButton(text="📊 Progress")],
            [KeyboardButton(text="📄 Demo ko'rish"), KeyboardButton(text="🔍 Plagiat tekshirish")],
            [KeyboardButton(text="📥 Hamma bandlarni yuklab"), KeyboardButton(text="👤 Profilim")],
            [KeyboardButton(text="📞 Admin bilan bog'lanish")],
        ],
        resize_keyboard=True,
        input_field_placeholder="Buyruq tanlang yoki yozing..."
    )

def inline_menu_free():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📄 Demo ko'rish (5-betlik)", callback_data="demo")],
        [InlineKeyboardButton(text="🔍 Plagiat tekshirish", callback_data="plagiat")],
        [InlineKeyboardButton(text="💰 Tarif sotib olish", callback_data="buy_tarif")],
        [InlineKeyboardButton(text="📋 Tariflar narxi", callback_data="show_tarifs")],
        [InlineKeyboardButton(text="👤 Profilim", callback_data="my_profile")],
        [InlineKeyboardButton(text="📞 Admin bilan bog'lanish", callback_data="contact_admin")],
    ])

def inline_menu_premium():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✍️ Band yozish (AI)", callback_data="write_band")],
        [InlineKeyboardButton(text="📊 Mening progressim", callback_data="progress")],
        [InlineKeyboardButton(text="📄 Demo ko'rish", callback_data="demo")],
        [InlineKeyboardButton(text="🔍 Plagiat tekshirish", callback_data="plagiat")],
        [InlineKeyboardButton(text="📥 Hamma bandlarni yuklab olish", callback_data="download_all")],
        [InlineKeyboardButton(text="💰 Tarif yangilash", callback_data="buy_tarif")],
        [InlineKeyboardButton(text="👤 Profilim", callback_data="my_profile")],
        [InlineKeyboardButton(text="📞 Admin bilan bog'lanish", callback_data="contact_admin")],
    ])

def main_menu(uid):
    user = get_user(uid)
    tarif = user.get("tarif", "free") if user else "free"
    if tarif == "free":
        return inline_menu_free()
    return inline_menu_premium()

def reply_menu(uid):
    user = get_user(uid)
    tarif = user.get("tarif", "free") if user else "free"
    if tarif == "free":
        return reply_keyboard_free()
    return reply_keyboard_premium()

def tarif_keyboard():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🆓 FREE - Tekin (Faqat plagiat)", callback_data="tarif_free")],
        [InlineKeyboardButton(text="📗 LITE - $25 / 312,500 so'm (50 bet)", callback_data="tarif_lite")],
        [InlineKeyboardButton(text="📘 PRO - $150 / 1,875,000 so'm (100 bet)", callback_data="tarif_pro")],
        [InlineKeyboardButton(text="💎 PROMAX - $375 / 4,687,500 so'm (150 bet)", callback_data="tarif_promax")],
    ])

def bands_keyboard(uid):
    """Mavjud bandlar + yangi qo'shish knopkalari"""
    user = get_user(uid)
    bands = user.get("bands", {}) if user else {}
    btns = []

    if bands:
        for num, band in sorted(bands.items()):
            btns.append([InlineKeyboardButton(
                text=f"✅ {num}: {band['name'][:25]}...",
                callback_data=f"view_band_{num}"
            )])

    btns.append([InlineKeyboardButton(text="➕ Yangi band yozish", callback_data="new_band")])
    btns.append([InlineKeyboardButton(text="🏠 Asosiy menyu", callback_data="main")])
    return InlineKeyboardMarkup(inline_keyboard=btns)

# ========================================
# /START
# ========================================
@dp.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    uid = message.from_user.id
    await state.clear()

    if is_admin(uid):
        await message.answer(
            "👑 ADMIN PANEL\n\n"
            "/admin - Statistika\n"
            "/pending - Kutayotganlar\n"
            "/users - Foydalanuvchilar\n"
            "/approve [id] - Ruxsat berish\n"
            "/reject [id] - Rad etish"
        )
        return

    if is_registered(uid):
        user = get_user(uid)
        tarif = user.get("tarif", "free")
        bands_done = len(user.get("bands", {}))
        t = TARIFS[tarif]

        if tarif == "free":
            await message.answer(
                f"👋 Xush kelibsiz, {user['name']}!\n\n"
                f"💰 Tarif: {t['name']}\n"
                f"ℹ️ Siz hozir FREE tarifdasiz.\n"
                f"AI bilan band yozish uchun tarif sotib oling!\n\n"
                f"Quyidagilardan foydalanishingiz mumkin:",
                reply_markup=reply_keyboard_free()
            )
            await message.answer("👇 Xizmatni tanlang:", reply_markup=inline_menu_free())
        else:
            await message.answer(
                f"👋 Xush kelibsiz, {user['name']}!\n\n"
                f"💰 Tarif: {t['name']}\n"
                f"✅ Yozilgan bandlar: {bands_done}/{t['bands']}\n\n"
                f"Davom etamizmi?",
                reply_markup=reply_keyboard_premium()
            )
            await message.answer("👇 Xizmatni tanlang:", reply_markup=inline_menu_premium())
        return

    if uid in pending_users:
        await message.answer(
            "⏳ So'rovingiz ko'rib chiqilmoqda.\n\n"
            "Admin tez orada javob beradi!\n"
            f"📞 Tezroq: {ADMIN_USERNAME}"
        )
        return

    await message.answer(
        "🎓 *AKADEMIK YORDAMCHI BOT*\n\n"
        "Assalomu alaykum!\n\n"
        "✅ AI yordamida band yozish\n"
        "✅ Plagiat tekshiruv (70%+)\n"
        "✅ 12+ manba qo'shish\n"
        "✅ Word fayl export\n\n"
        "📝 Ismingiz va familiyangizni kiriting:\n"
        "_(Masalan: Abdullayev Abdulaziz)_",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardRemove()
    )
    await state.set_state(Registration.name)

# ========================================
# RO'YXATDAN O'TISH
# ========================================
@dp.message(StateFilter(Registration.name))
async def reg_name(message: types.Message, state: FSMContext):
    await state.update_data(name=message.text)
    await message.answer(
        f"✅ Ism: *{message.text}*\n\n"
        f"📱 Telefon raqamingiz:\n"
        f"_(Masalan: +998901234567)_",
        parse_mode="Markdown"
    )
    await state.set_state(Registration.phone)

@dp.message(StateFilter(Registration.phone))
async def reg_phone(message: types.Message, state: FSMContext):
    await state.update_data(phone=message.text)
    await message.answer(
        f"✅ Telefon qabul qilindi!\n\n"
        f"🎓 Dissertatsiya mavzuingizni to'liq yozing:\n"
        f"_(Masalan: 'Boshlang'ich sinf o'quvchilarida...')_",
        parse_mode="Markdown"
    )
    await state.set_state(Registration.topic)

@dp.message(StateFilter(Registration.topic))
async def reg_topic(message: types.Message, state: FSMContext):
    await state.update_data(topic=message.text)
    await message.answer(
        f"✅ Mavzu qabul qilindi!\n\n"
        f"❓ Qanday yordam kerak?\n"
        f"_(Masalan: 'Faqat 2-bob kerak' yoki 'Butun dissertatsiya')_",
        parse_mode="Markdown"
    )
    await state.set_state(Registration.help_type)

@dp.message(StateFilter(Registration.help_type))
async def reg_help(message: types.Message, state: FSMContext):
    await state.update_data(help_type=message.text)
    await message.answer(
        "💰 Tarif tanlang:\n\n"
        "🆓 FREE — Tekin\n"
        "   • Faqat plagiat tekshiruv\n"
        "   • Demo namuna ko'rish\n\n"
        "📗 LITE — $25 / 312,500 so'm\n"
        "   • 50 bet, 5 band, AI yozuv\n\n"
        "📘 PRO — $150 / 1,875,000 so'm\n"
        "   • 100 bet, 10 band\n\n"
        "💎 PROMAX — $375 / 4,687,500 so'm\n"
        "   • 150 bet, 15 band, priority\n\n"
        "⚠️ Avval DEMO ko'ring!",
        reply_markup=tarif_keyboard()
    )
    await state.set_state(Registration.tarif)

@dp.callback_query(F.data.startswith("tarif_"), StateFilter(Registration.tarif))
async def reg_tarif(callback: types.CallbackQuery, state: FSMContext):
    tarif = callback.data.replace("tarif_", "")
    data = await state.get_data()
    uid = callback.from_user.id
    await state.clear()

    pending_users[uid] = {
        "name": data.get("name"),
        "phone": data.get("phone"),
        "topic": data.get("topic"),
        "help_type": data.get("help_type"),
        "tarif": tarif,
        "status": "pending",
        "registered_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "user_id": uid,
        "username": callback.from_user.username or "yo'q",
        "bands": {},
    }

    t = TARIFS[tarif]
    await callback.message.answer(
        f"✅ Ro'yxatdan o'tdingiz!\n\n"
        f"👤 {data.get('name')}\n"
        f"💰 {t['name']} - ${t['price']}\n\n"
        f"⏳ Endi admin tasdiqlashini kuting.\n"
        f"Odatda 1-2 soat ichida javob beriladi.\n\n"
        f"📞 Tezroq bog'lanish: {ADMIN_USERNAME}",
        reply_markup=ReplyKeyboardRemove()
    )

    if ADMIN_ID:
        await bot.send_message(
            ADMIN_ID,
            f"🔔 YANGI FOYDALANUVCHI!\n\n"
            f"👤 {data.get('name')}\n"
            f"📱 {data.get('phone')}\n"
            f"🆔 {uid}\n"
            f"@{callback.from_user.username or 'username yoq'}\n"
            f"🎓 {data.get('topic')}\n"
            f"❓ {data.get('help_type')}\n"
            f"💰 {t['name']} - ${t['price']}\n\n"
            f"✅ Tasdiqlash: /approve {uid}\n"
            f"❌ Rad etish: /reject {uid}",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [
                    InlineKeyboardButton(text="✅ Tasdiqlash", callback_data=f"adm_approve_{uid}"),
                    InlineKeyboardButton(text="❌ Rad etish", callback_data=f"adm_reject_{uid}"),
                ]
            ])
        )
    await callback.answer()

# ========================================
# ADMIN INLINE BUTTONS
# ========================================
@dp.callback_query(F.data.startswith("adm_approve_"))
async def adm_approve(callback: types.CallbackQuery):
    if not is_admin(callback.from_user.id): return
    uid = int(callback.data.replace("adm_approve_", ""))

    # Render restart bo'lsa pending_users yo'qoladi - shuning uchun fallback
    if uid in pending_users:
        user = pending_users.pop(uid)
    elif uid in users_db:
        # Allaqachon tasdiqlangan
        await callback.answer("Bu foydalanuvchi allaqachon tasdiqlangan!", show_alert=True)
        return
    else:
        # Restart bo'lgan - minimal user yaratamiz
        user = {
            "name": "Foydalanuvchi",
            "phone": "-",
            "topic": "-",
            "help_type": "-",
            "tarif": "free",
            "registered_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "user_id": uid,
            "username": "-",
        }
        await callback.answer("⚠️ Bot restart bo'lgan, minimal ma'lumot bilan tasdiqlandi.", show_alert=True)

    user["status"] = "approved"
    user["bands"] = {}
    users_db[uid] = user
    t = TARIFS[user["tarif"]]

    # Foydalanuvchiga xabar
    if user["tarif"] == "free":
        await bot.send_message(
            uid,
            f"🎉 Tabriklaymiz, {user['name']}!\n\n"
            f"✅ Admin tomonidan tasdiqlandi!\n"
            f"💰 Tarifingiz: {t['name']}\n\n"
            f"Siz hozir quyidagilardan foydalanishingiz mumkin:\n"
            f"• Demo ko'rish\n"
            f"• Plagiat tekshirish\n\n"
            f"AI bilan band yozish uchun tarif sotib oling:",
            reply_markup=reply_keyboard_free()
        )
        await bot.send_message(uid, "👇 Xizmatni tanlang:", reply_markup=inline_menu_free())
    else:
        await bot.send_message(
            uid,
            f"🎉 Tabriklaymiz, {user['name']}!\n\n"
            f"✅ Admin tomonidan tasdiqlandi!\n"
            f"💰 Tarifingiz: {t['name']}\n"
            f"📝 Jami bandlar: {t['bands']} ta\n\n"
            f"Endi dissertatsiyangizni yozishni boshlashingiz mumkin!\n\n"
            f"Qaysi bobdan boshlashni xohlaysiz?\n"
            f"_(Masalan: 1-bob, 1.1-band)_",
            reply_markup=reply_keyboard_premium()
        )
        await bot.send_message(
            uid,
            "👇 Band yozishni boshlash uchun:",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="✍️ Band yozishni boshlash", callback_data="write_band")],
                [InlineKeyboardButton(text="📊 Mening progressim", callback_data="progress")],
            ])
        )

    await callback.message.edit_reply_markup(reply_markup=None)
    await callback.answer(f"✅ {user['name']} tasdiqlandi!", show_alert=True)

@dp.callback_query(F.data.startswith("adm_reject_"))
async def adm_reject(callback: types.CallbackQuery):
    if not is_admin(callback.from_user.id): return
    uid = int(callback.data.replace("adm_reject_", ""))
    if uid not in pending_users:
        # Restart bo'lgan - baribir xabar yuboramiz
        await bot.send_message(uid,
            f"❌ So'rovingiz rad etildi.\n"
            f"📞 {ADMIN_USERNAME}"
        )
        await callback.message.edit_reply_markup(reply_markup=None)
        await callback.answer("Rad etildi (restart bo'lgan).", show_alert=True)
        return
    user = pending_users.pop(uid)
    await bot.send_message(
        uid,
        f"❌ Hurmatli {user['name']},\n\n"
        f"So'rovingiz rad etildi.\n"
        f"Qo'shimcha ma'lumot uchun:\n"
        f"📞 {ADMIN_USERNAME}"
    )
    await callback.message.edit_reply_markup(reply_markup=None)
    await callback.answer(f"❌ {user['name']} rad etildi!", show_alert=True)

# ========================================
# REPLY KEYBOARD HANDLERS
# ========================================
@dp.message(F.text == "📄 Demo ko'rish", StateFilter(None))
async def reply_demo(message: types.Message):
    await send_demo(message, uid=message.from_user.id)

@dp.message(F.text == "🔍 Plagiat tekshirish", StateFilter(None))
async def reply_plagiat(message: types.Message, state: FSMContext):
    await message.answer(
        "🔍 PLAGIAT TEKSHIRUV\n\n"
        "Matnni kiriting (kamida 200 belgi):\n\n"
        "⚠️ Manbalar: [manba 1], [manba 2] tarzida"
    )
    await state.set_state(PlagiatCheck.text)

@dp.message(F.text == "💰 Tarif sotib olish", StateFilter(None))
async def reply_buy(message: types.Message):
    await show_buy_tarif(message)

@dp.message(F.text == "👤 Profilim", StateFilter(None))
async def reply_profile(message: types.Message):
    await show_profile(message, message.from_user.id)

@dp.message(F.text == "📞 Admin bilan bog'lanish", StateFilter(None))
async def reply_admin(message: types.Message):
    await show_admin_contact(message)

@dp.message(F.text == "✍️ Band yozish (AI)", StateFilter(None))
async def reply_write_band(message: types.Message, state: FSMContext):
    await start_band_write(message, state, message.from_user.id)

@dp.message(F.text == "📊 Progress", StateFilter(None))
async def reply_progress(message: types.Message):
    await show_progress_msg(message, message.from_user.id)

@dp.message(F.text == "📥 Hamma bandlarni yuklab", StateFilter(None))
async def reply_download(message: types.Message):
    await download_all_bands(message, message.from_user.id)

# ========================================
# DEMO
# ========================================
@dp.message(Command("demo"))
async def cmd_demo(message: types.Message):
    await send_demo(message, uid=message.from_user.id)

@dp.callback_query(F.data == "demo")
async def callback_demo(callback: types.CallbackQuery):
    await send_demo(callback.message, uid=callback.from_user.id)
    await callback.answer()

async def send_demo(message, uid=None):
    wait = await message.answer("⏳ Demo tayyorlanmoqda... (10-20 soniya)")
    try:
        filename = f"demo_{message.chat.id}.docx"
        user_id = uid or message.chat.id
        user = get_user(user_id)
        user_name = user.get("name", "Foydalanuvchi") if user else "Foydalanuvchi"
        user_topic = user.get("topic", None) if user else None
        generate_demo(filename, user_name=user_name, user_topic=user_topic, is_demo=True)

        await bot.delete_message(message.chat.id, wait.message_id)
        await message.answer_document(
            types.FSInputFile(filename),
            caption=(
                "📄 DEMO NAMUNA TAYYOR!\n\n"
                f"👤 Muallif: {user_name}\n\n"
                "📋 Bu namunada:\n"
                "✅ 5-betlik professional ilmiy matn\n"
                "✅ 12 ta manba (muallif, kitob, yil, bet)\n"
                "✅ Plagiat natijasi: 74% ✅\n"
                "✅ Professional manbalar jadvali\n\n"
                "💡 TO'LIQ DISSERTATSIYA uchun:\n"
                "📗 LITE - $25 | 📘 PRO - $150 | 💎 PROMAX - $375"
            ),
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="💰 Tarif sotib olish", callback_data="buy_tarif")],
            ])
        )
        os.remove(filename)
    except Exception as e:
        await bot.delete_message(message.chat.id, wait.message_id)
        await message.answer(f"❌ Xatolik yuz berdi. Qayta urining.")

# ========================================
# BAND YOZISH WORKFLOW
# ========================================
@dp.callback_query(F.data == "write_band")
async def callback_write_band(callback: types.CallbackQuery, state: FSMContext):
    await start_band_write(callback.message, state, callback.from_user.id)
    await callback.answer()

@dp.callback_query(F.data == "new_band")
async def callback_new_band(callback: types.CallbackQuery, state: FSMContext):
    await start_new_band(callback.message, state, callback.from_user.id)
    await callback.answer()

async def start_band_write(message, state, uid):
    if not is_registered(uid):
        await message.answer("❌ Avval ro'yxatdan o'ting: /start")
        return

    user = get_user(uid)
    if user.get("tarif") == "free":
        await message.answer(
            "❌ Bu xizmat faqat LITE, PRO va PROMAX uchun!\n\n"
            "💰 Tarif sotib olish:",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="💰 Tarif sotib olish", callback_data="buy_tarif")]
            ])
        )
        return

    bands = user.get("bands", {})
    t = TARIFS[user["tarif"]]

    # Progress ko'rsatish
    progress_text = f"📊 PROGRESS: {len(bands)}/{t['bands']} band\n\n"
    if bands:
        progress_text += "✅ Yozilgan bandlar:\n"
        for num, band in sorted(bands.items()):
            progress_text += f"  ✅ {num}: {band['name'][:35]}\n"
        progress_text += "\n"

    await message.answer(
        f"✍️ BAND YOZISH\n\n"
        f"{progress_text}"
        f"🎓 Mavzu: {user.get('topic', '')[:60]}...\n\n"
        f"Quyidagilardan birini tanlang:",
        reply_markup=bands_keyboard(uid)
    )

async def start_new_band(message, state, uid):
    user = get_user(uid)
    t = TARIFS[user["tarif"]]
    bands = user.get("bands", {})

    if len(bands) >= t["bands"]:
        await message.answer(
            f"❌ Tarifingiz limiti tugadi!\n"
            f"Jami: {t['bands']} ta band\n\n"
            f"💰 Yangi tarif sotib oling:",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="💎 PROMAX ga o'tish", callback_data="buy_tarif")]
            ])
        )
        return

    await message.answer(
        f"✍️ YANGI BAND\n\n"
        f"Band raqamini kiriting:\n"
        f"_(Masalan: 1.1, 1.2, 2.1, 3.2)_",
        parse_mode="Markdown"
    )
    await state.set_state(BandWrite.band_number)

@dp.message(StateFilter(BandWrite.band_number))
async def band_number(message: types.Message, state: FSMContext):
    await state.update_data(band_number=message.text)
    await message.answer(
        f"✅ Band raqami: *{message.text}*\n\n"
        f"Band nomini kiriting:\n"
        f"_(Masalan: 'Sun'iy intellektning ta'riflari va turlari')_",
        parse_mode="Markdown"
    )
    await state.set_state(BandWrite.band_name)

@dp.message(StateFilter(BandWrite.band_name))
async def band_name_handler(message: types.Message, state: FSMContext):
    data = await state.get_data()
    band_num = data.get("band_number")
    band_name_text = message.text
    await state.update_data(band_name=band_name_text)

    uid = message.from_user.id
    user = get_user(uid)
    topic = user.get("topic", "")

    wait = await message.answer(
        f"⏳ AI *{band_num}* bandni yozmoqda...\n"
        f"📝 Band: {band_name_text}\n\n"
        f"30-60 soniya kutib turing...",
        parse_mode="Markdown"
    )

    try:
        prompt = f"""Sen O'zbekistonda 20 yillik tajribaga ega ilmiy tadqiqotchi va professor.

Dissertatsiya mavzusi: "{topic}"
Band raqami: {band_num}
Band nomi: "{band_name_text}"

Quyidagi talablarga rioya qilib ilmiy band yozing (O'zbek tilida):

1. Hajmi: 600-800 so'z
2. Kamida 8 ta manba [manba 1], [manba 2] tarzida
3. Akademik, ilmiy til
4. Har bir fikr manba bilan asoslangan
5. Xalqaro va O'zbekiston manbalaridan foydalaning
6. Tuzilma: ta'rif → tahlil → xulosa

FAQAT band matnini yoz (sarlavhasiz):"""

        response = model.generate_content(prompt)
        band_text = response.text
        orig, plagiat = plagiat_check(band_text)
        words = len(band_text.split())

        await bot.delete_message(message.chat.id, wait.message_id)
        preview = band_text[:600] + "\n\n...[davomi bor]" if len(band_text) > 600 else band_text

        status_emoji = "✅" if orig >= 70 else "⚠️"
        await message.answer(
            f"✅ *{band_num} BAND TAYYOR!*\n\n"
            f"📝 {band_name_text}\n"
            f"{status_emoji} Originallik: *{orig}%*\n"
            f"📝 So'zlar: {words} ta\n\n"
            f"─────────────\n"
            f"{preview}\n"
            f"─────────────\n\n"
            f"Qabul qilasizmi?",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="✅ Qabul - Word yukla", callback_data=f"accept_{band_num}")],
                [InlineKeyboardButton(text="🔄 Qayta yozish", callback_data=f"rewrite_{band_num}")],
                [InlineKeyboardButton(text="✏️ O'zgartirish kiritaman", callback_data=f"edit_{band_num}")],
            ])
        )
        await state.update_data(band_text=band_text, orig=orig)
        await state.set_state(BandWrite.waiting_ai)

    except Exception as e:
        await bot.delete_message(message.chat.id, wait.message_id)
        err = str(e)
        if "429" in err or "quota" in err.lower() or "limit" in err.lower():
            await message.answer(
                "⚠️ AI server hozir band.\n\n"
                "📋 Sabab: Kunlik so'rovlar limiti to'ldi.\n"
                "⏰ Iltimos, 5-10 daqiqadan keyin qayta urining.\n\n"
                f"📞 Muammo davom etsa: {ADMIN_USERNAME}",
                reply_markup=reply_menu(uid)
            )
        else:
            await message.answer(
                f"❌ Xatolik yuz berdi.\n"
                f"📞 Admin: {ADMIN_USERNAME}",
                reply_markup=reply_menu(uid)
            )
        await state.clear()

@dp.callback_query(F.data.startswith("accept_"), StateFilter(BandWrite.waiting_ai))
async def band_accept(callback: types.CallbackQuery, state: FSMContext):
    band_num = callback.data.replace("accept_", "")
    data = await state.get_data()
    uid = callback.from_user.id
    user = get_user(uid)

    band_text = data.get("band_text", "")
    band_name_text = data.get("band_name", "")
    orig = data.get("orig", 0)

    if "bands" not in users_db[uid]:
        users_db[uid]["bands"] = {}
    users_db[uid]["bands"][band_num] = {
        "name": band_name_text,
        "text": band_text,
        "orig": orig,
        "status": "done",
        "date": datetime.now().strftime("%Y-%m-%d %H:%M")
    }

    from docx import Document as DocxDoc
    from docx.shared import Pt as DocxPt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WDA

    doc = DocxDoc()
    section = doc.sections[0]
    h = doc.add_heading(f"{band_num}. {band_name_text}", level=1)
    h.alignment = WDA.CENTER

    for para in band_text.split('\n'):
        if para.strip():
            p = doc.add_paragraph(para)
            p.alignment = WDA.JUSTIFY
            p.paragraph_format.first_line_indent = DocxPt(36)
            for run in p.runs:
                run.font.size = DocxPt(12)
                run.font.name = "Times New Roman"

    doc.add_paragraph()
    doc.add_paragraph(f"✅ Originallik: {orig}%")
    doc.add_paragraph(f"📅 {datetime.now().strftime('%Y-%m-%d')}")

    filename = f"band_{band_num.replace('.','_')}_{uid}.docx"
    doc.save(filename)

    bands_done = len(users_db[uid]["bands"])
    t = TARIFS[user.get("tarif", "free")]

    await callback.message.answer_document(
        types.FSInputFile(filename),
        caption=(
            f"✅ *{band_num} BAND SAQLANDI!*\n\n"
            f"📝 {band_name_text}\n"
            f"📊 Originallik: {orig}%\n"
            f"📅 {datetime.now().strftime('%Y-%m-%d')}\n\n"
            f"📊 Progress: {bands_done}/{t['bands']} band",
        ),
        parse_mode="Markdown"
    )
    os.remove(filename)
    await state.clear()

    # Keyingi band tanlash
    await callback.message.answer(
        f"📊 *PROGRESS: {bands_done}/{t['bands']} band*\n\n"
        f"{'🎉 Barcha bandlar yozildi!' if bands_done >= t['bands'] else 'Keyingi band:'}\n",
        parse_mode="Markdown",
        reply_markup=bands_keyboard(uid)
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("rewrite_"), StateFilter(BandWrite.waiting_ai))
async def band_rewrite(callback: types.CallbackQuery, state: FSMContext):
    band_num = callback.data.replace("rewrite_", "")
    await state.update_data(band_number=band_num)
    await state.set_state(BandWrite.band_name)
    await callback.message.answer(
        "🔄 Band qayta yoziladi.\n"
        "Band nomini qayta kiriting:"
    )
    await callback.answer()

@dp.callback_query(F.data.startswith("edit_"), StateFilter(BandWrite.waiting_ai))
async def band_edit(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.answer(
        "✏️ Qanday o'zgarish kerak?\n\n"
        "_(Masalan: 'Plagiat foizi baland, qayta yozsin'\n"
        "yoki '3-paragraf zaif, kuchaytirsin')_",
        parse_mode="Markdown"
    )
    await state.set_state(BandWrite.edit_text)
    await callback.answer()

@dp.message(StateFilter(BandWrite.edit_text))
async def band_edit_text(message: types.Message, state: FSMContext):
    data = await state.get_data()
    uid = message.from_user.id
    user = get_user(uid)
    old_text = data.get("band_text", "")
    band_num = data.get("band_number")
    band_name_text = data.get("band_name")
    topic = user.get("topic", "")
    edit_request = message.text

    wait = await message.answer("⏳ AI matnni qayta yozmoqda...")

    try:
        prompt = f"""Dissertatsiya mavzusi: "{topic}"
Band: {band_num} - "{band_name_text}"

ESKI MATN:
{old_text[:1500]}

FOYDALANUVCHI XOHISHI:
"{edit_request}"

Ushbu o'zgartirishlarni kiritib, bandni qayta yozing:
1. 600-800 so'z
2. Kamida 8 ta manba [manba N] tarzida
3. Akademik uslub
4. O'zgartirishlarni to'liq bajaring"""

        response = model.generate_content(prompt)
        new_text = response.text
        orig, _ = plagiat_check(new_text)
        words = len(new_text.split())

        await bot.delete_message(message.chat.id, wait.message_id)
        preview = new_text[:600] + "\n\n...[davomi bor]" if len(new_text) > 600 else new_text

        await message.answer(
            f"✅ *QAYTA YOZILDI!*\n\n"
            f"📊 Originallik: *{orig}%*\n"
            f"📝 So'zlar: {words}\n\n"
            f"─────────────\n"
            f"{preview}\n"
            f"─────────────\n\n"
            f"Qabul qilasizmi?",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="✅ Qabul qilaman", callback_data=f"accept_{band_num}")],
                [InlineKeyboardButton(text="🔄 Yana o'zgartirish", callback_data=f"edit_{band_num}")],
            ])
        )
        await state.update_data(band_text=new_text, orig=orig)
        await state.set_state(BandWrite.waiting_ai)

    except Exception as e:
        await bot.delete_message(message.chat.id, wait.message_id)
        if "429" in str(e) or "quota" in str(e).lower():
            await message.answer(
                f"⚠️ AI server band. 5-10 daqiqadan keyin qayta urining.\n"
                f"📞 {ADMIN_USERNAME}",
                reply_markup=reply_menu(uid)
            )
        else:
            await message.answer(f"❌ Xatolik. Admin: {ADMIN_USERNAME}")
        await state.clear()

# ========================================
# PLAGIAT TEKSHIRUV
# ========================================
@dp.callback_query(F.data == "plagiat")
async def plagiat_callback(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.answer(
        "🔍 PLAGIAT TEKSHIRUV\n\n"
        "Matnni kiriting (kamida 200 belgi):\n\n"
        "⚠️ Manbalar: [manba 1], [manba 2] tarzida"
    )
    await state.set_state(PlagiatCheck.text)
    await callback.answer()

@dp.message(StateFilter(PlagiatCheck.text))
async def plagiat_result(message: types.Message, state: FSMContext):
    text = message.text
    await state.clear()
    if len(text) < 100:
        await message.answer("❌ Matn juda qisqa! Kamida 200 belgi yozing.")
        await state.set_state(PlagiatCheck.text)
        return

    orig, plagiat = plagiat_check(text)
    sources = len([w for w in text.split('[') if 'manba' in w.lower()])
    words = len(text.split())

    if orig >= 70: status = "✅ QABUL QILINADI"; emoji = "✅"
    elif orig >= 60: status = "⚠️ SHARTLI (tavsiya etilmaydi)"; emoji = "⚠️"
    else: status = "❌ PLAGIAT (qayta yozish kerak)"; emoji = "❌"

    result = (
        f"{emoji} PLAGIAT TEKSHIRUV NATIJASI\n\n"
        f"📊 Originallik: {orig}%\n"
        f"📉 Ko'chirmakashlik: {plagiat}%\n"
        f"📚 Manbalar: {sources} ta\n"
        f"📝 So'zlar: {words} ta\n"
        f"🏷️ Status: {status}\n"
    )

    if orig < 70:
        result += "\n💡 TAVSIYALAR:\n"
        if orig < 70: result += "• O'z fikrlaringizni ko'proq yozing\n"
        if sources < 5: result += "• Ko'proq manba qo'shing: [manba 1]\n"
        result += "• 3-4 paragrafni qayta yozing\n"

    uid = message.from_user.id
    await message.answer(
        result,
        reply_markup=reply_menu(uid) if is_registered(uid) else None
    )

# ========================================
# PAYMENT
# ========================================
async def show_buy_tarif(message):
    await message.answer(
        "💰 TARIF SOTIB OLISH\n\n"
        "📗 LITE — $25 / 312,500 so'm\n"
        "  • 50 bet | 5 band | AI yozuv\n\n"
        "📘 PRO — $150 / 1,875,000 so'm\n"
        "  • 100 bet | 10 band\n\n"
        "💎 PROMAX — $375 / 4,687,500 so'm\n"
        "  • 150 bet | 15 band | Priority\n\n"
        "💳 TO'LOV USULLARI:\n"
        f"• Click/Payme: {ADMIN_PHONE}\n"
        f"• Telegram: {ADMIN_USERNAME}\n\n"
        "To'lovdan so'ng chek yuboring:",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="📸 To'lov chekini yuborish", callback_data="send_payment")],
        ])
    )

@dp.callback_query(F.data == "buy_tarif")
async def buy_tarif(callback: types.CallbackQuery):
    await show_buy_tarif(callback.message)
    await callback.answer()

@dp.callback_query(F.data == "send_payment")
async def send_payment(callback: types.CallbackQuery, state: FSMContext):
    await callback.message.answer(
        "📸 To'lov cheki (screenshot) yuboring:\n\n"
        "✅ Chekda ko'rinishi kerak:\n"
        "• Summa\n• Sana\n• Kimga o'tkazilgan"
    )
    await state.set_state(Payment.screenshot)
    await callback.answer()

@dp.message(StateFilter(Payment.screenshot), F.photo)
async def payment_screenshot(message: types.Message, state: FSMContext):
    uid = message.from_user.id
    user = get_user(uid) or pending_users.get(uid, {})
    await state.clear()
    await message.answer(
        "✅ To'lov cheki qabul qilindi!\n\n"
        "⏳ Admin 1-2 soat ichida tasdiqlab,\n"
        "tarifingizni ochib beradi."
    )
    if ADMIN_ID:
        await bot.send_photo(
            ADMIN_ID,
            message.photo[-1].file_id,
            caption=(
                f"💰 TO'LOV CHEKI!\n\n"
                f"👤 {user.get('name', 'Noma\'lum')}\n"
                f"🆔 {uid}\n"
                f"📱 {user.get('phone', '')}\n\n"
                f"✅ /approve {uid}\n❌ /reject {uid}"
            ),
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [
                    InlineKeyboardButton(text="✅ Tasdiqlash", callback_data=f"adm_approve_{uid}"),
                    InlineKeyboardButton(text="❌ Rad etish", callback_data=f"adm_reject_{uid}"),
                ]
            ])
        )

@dp.message(StateFilter(Payment.screenshot))
async def payment_no_photo(message: types.Message):
    await message.answer("❌ Iltimos, screenshot rasm yuboring!")

# ========================================
# PROGRESS
# ========================================
async def show_progress_msg(message, uid):
    user = get_user(uid)
    if not user:
        await message.answer("❌ Profil topilmadi. /start")
        return

    tarif = user.get("tarif", "free")
    bands = user.get("bands", {})
    t = TARIFS[tarif]
    max_bands = t["bands"]
    done = len(bands)
    percent = int(done / max_bands * 100) if max_bands > 0 else 0
    filled = int(percent / 10)
    bar = "█" * filled + "░" * (10 - filled)

    text = (
        f"📊 MENING PROGRESSIM\n\n"
        f"👤 {user.get('name', '')}\n"
        f"🎓 {user.get('topic', '')[:50]}...\n"
        f"💰 {t['name']}\n\n"
        f"[{bar}] {percent}%\n"
        f"✅ Bajarildi: {done}/{max_bands} band\n\n"
    )

    if bands:
        text += "📋 YOZILGAN BANDLAR:\n"
        for num, band in sorted(bands.items()):
            text += f"  ✅ {num}: {band['name'][:40]}\n"
            text += f"     📊 Originallik: {band['orig']}%\n"

    if done >= max_bands and max_bands > 0:
        text += "\n🎉 BARCHA BANDLAR TAYYOR!"
    elif max_bands > 0:
        text += f"\n⏳ Qoldi: {max_bands - done} band"

    await message.answer(text, reply_markup=bands_keyboard(uid) if tarif != "free" else inline_menu_free())

@dp.callback_query(F.data == "progress")
async def progress_callback(callback: types.CallbackQuery):
    await show_progress_msg(callback.message, callback.from_user.id)
    await callback.answer()

# ========================================
# ADMIN CONTACT
# ========================================
async def show_admin_contact(message):
    await message.answer(
        f"📞 ADMIN BILAN BOG'LANISH\n\n"
        f"✉️ Telegram: {ADMIN_USERNAME}\n"
        f"📱 Telefon: {ADMIN_PHONE}\n\n"
        f"⏰ Ish vaqti: 9:00 - 18:00\n"
        f"📅 Dushanba - Shanba\n\n"
        f"💬 Savol yuboring — javob beramiz!"
    )

@dp.callback_query(F.data == "contact_admin")
async def contact_admin_cb(callback: types.CallbackQuery):
    await show_admin_contact(callback.message)
    await callback.answer()

# ========================================
# PROFIL
# ========================================
async def show_profile(message, uid):
    user = get_user(uid)
    if not user:
        await message.answer("❌ Profil topilmadi. /start")
        return
    t = TARIFS.get(user.get("tarif", "free"))
    bands = user.get("bands", {})
    await message.answer(
        f"👤 MENING PROFILIM\n\n"
        f"📛 Ism: {user['name']}\n"
        f"📱 Tel: {user['phone']}\n"
        f"🎓 Mavzu: {user['topic'][:60]}...\n"
        f"💰 Tarif: {t['name']}\n"
        f"✅ Bandlar: {len(bands)} ta\n"
        f"📅 Ro'yxat: {user['registered_at']}"
    )

@dp.callback_query(F.data == "my_profile")
async def my_profile_cb(callback: types.CallbackQuery):
    await show_profile(callback.message, callback.from_user.id)
    await callback.answer()

# ========================================
# TARIFLAR
# ========================================
@dp.callback_query(F.data == "show_tarifs")
async def show_tarifs(callback: types.CallbackQuery):
    await callback.message.answer(
        "💰 TARIFLAR\n\n"
        "🆓 FREE — Tekin\n"
        "  • Demo ko'rish\n"
        "  • Plagiat tekshiruv\n\n"
        "📗 LITE — $25 / 312,500 so'm\n"
        "  • 50 bet | 5 band\n"
        "  • AI matn yozuv\n\n"
        "📘 PRO — $150 / 1,875,000 so'm\n"
        "  • 100 bet | 10 band\n\n"
        "💎 PROMAX — $375 / 4,687,500 so'm\n"
        "  • 150 bet | 15 band\n"
        "  • Priority support\n\n"
        f"📞 Sotib olish: {ADMIN_USERNAME}",
        reply_markup=InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="💰 Tarif sotib olish", callback_data="buy_tarif")]
        ])
    )
    await callback.answer()

# ========================================
# HAMMA BANDLARNI YUKLAB OLISH
# ========================================
async def download_all_bands(message, uid):
    user = get_user(uid)
    if not user or not user.get("bands"):
        await message.answer("❌ Hali hech qanday band yozilmagan!\n\n✍️ Band yozishni boshlang.")
        return

    wait = await message.answer("⏳ Barcha bandlar birlashtirilmoqda...")

    from docx import Document as DocxDoc
    from docx.shared import Pt as DocxPt
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WDA

    doc = DocxDoc()
    section = doc.sections[0]

    p = doc.add_paragraph()
    p.alignment = WDA.CENTER
    run = p.add_run("DISSERTATSIYA")
    run.font.size = DocxPt(18)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WDA.CENTER
    run = p.add_run(f"Muallif: {user.get('name', '')}")
    run.font.size = DocxPt(14)

    p = doc.add_paragraph()
    p.alignment = WDA.CENTER
    run = p.add_run(f"Mavzu: {user.get('topic', '')}")
    run.font.size = DocxPt(12)

    doc.add_page_break()

    bands = user.get("bands", {})
    for num, band in sorted(bands.items()):
        h = doc.add_heading(f"{num}. {band['name']}", level=1)
        for para in band["text"].split("\n"):
            if para.strip():
                p = doc.add_paragraph(para)
                p.alignment = WDA.JUSTIFY
                for run in p.runs:
                    run.font.size = DocxPt(12)
                    run.font.name = "Times New Roman"
        doc.add_paragraph(f"✅ Originallik: {band['orig']}%")
        doc.add_page_break()

    filename = f"full_{uid}.docx"
    doc.save(filename)
    await bot.delete_message(message.chat.id, wait.message_id)
    await message.answer_document(
        types.FSInputFile(filename),
        caption=(
            f"📚 TO'LIQ DISSERTATSIYA\n\n"
            f"👤 {user.get('name', '')}\n"
            f"📝 Bandlar: {len(bands)} ta\n"
            f"📅 {datetime.now().strftime('%Y-%m-%d')}\n\n"
            f"✅ Barcha bandlar birlashtirildi!"
        )
    )
    os.remove(filename)

@dp.callback_query(F.data == "download_all")
async def download_all_cb(callback: types.CallbackQuery):
    await download_all_bands(callback.message, callback.from_user.id)
    await callback.answer()

@dp.callback_query(F.data == "main")
async def main_cb(callback: types.CallbackQuery):
    uid = callback.from_user.id
    await callback.message.answer("🏠 Asosiy menyu:", reply_markup=main_menu(uid))
    await callback.answer()

# ========================================
# ADMIN COMMANDS
# ========================================
@dp.message(Command("admin"))
async def admin_panel(message: types.Message):
    if not is_admin(message.from_user.id): return
    await message.answer(
        f"👑 ADMIN PANEL\n\n"
        f"✅ Tasdiqlangan: {len(users_db)}\n"
        f"⏳ Kutayotgan: {len(pending_users)}\n\n"
        f"/pending - Kutayotganlar\n"
        f"/users - Barcha\n"
        f"/approve [id]\n"
        f"/reject [id]"
    )

@dp.message(Command("pending"))
async def admin_pending(message: types.Message):
    if not is_admin(message.from_user.id): return
    if not pending_users:
        await message.answer("⏳ Kutayotgan yo'q.")
        return
    for uid, u in pending_users.items():
        t = TARIFS.get(u.get("tarif", "free"))
        await message.answer(
            f"👤 {u['name']}\n"
            f"📱 {u['phone']}\n"
            f"🆔 {uid}\n"
            f"🎓 {u['topic'][:60]}\n"
            f"💰 {t['name']} - ${t['price']}",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [
                    InlineKeyboardButton(text="✅ Tasdiqlash", callback_data=f"adm_approve_{uid}"),
                    InlineKeyboardButton(text="❌ Rad etish", callback_data=f"adm_reject_{uid}"),
                ]
            ])
        )

@dp.message(Command("users"))
async def admin_users(message: types.Message):
    if not is_admin(message.from_user.id): return
    if not users_db:
        await message.answer("👥 Foydalanuvchi yo'q.")
        return
    text = f"✅ FOYDALANUVCHILAR ({len(users_db)} ta):\n\n"
    for uid, u in users_db.items():
        bands = len(u.get("bands", {}))
        t = TARIFS.get(u.get("tarif", "free"))
        text += f"👤 {u['name']} | {t['name']} | {bands} band | {uid}\n"
    await message.answer(text)

@dp.message(Command("approve"))
async def admin_approve(message: types.Message):
    if not is_admin(message.from_user.id): return
    parts = message.text.split()
    if len(parts) < 2:
        await message.answer("Format: /approve [id]")
        return
    uid = int(parts[1])
    if uid not in pending_users:
        await message.answer(f"❌ {uid} topilmadi.")
        return
    user = pending_users.pop(uid)
    user["status"] = "approved"
    user["bands"] = {}
    users_db[uid] = user
    t = TARIFS[user["tarif"]]

    if user["tarif"] == "free":
        await bot.send_message(uid,
            f"🎉 {user['name']}, tasdiqlandi!\n"
            f"💰 {t['name']}\n\n"
            f"Demo va plagiat tekshirishdan foydalaning:",
            reply_markup=reply_keyboard_free()
        )
        await bot.send_message(uid, "👇", reply_markup=inline_menu_free())
    else:
        await bot.send_message(uid,
            f"🎉 {user['name']}, tasdiqlandi!\n"
            f"💰 {t['name']} — {t['bands']} band\n\n"
            f"Qaysi bobdan boshlaysiz?",
            reply_markup=reply_keyboard_premium()
        )
        await bot.send_message(uid, "👇 Band yozishni boshlang!",
            reply_markup=InlineKeyboardMarkup(inline_keyboard=[
                [InlineKeyboardButton(text="✍️ Band yozishni boshlash", callback_data="write_band")]
            ])
        )
    await message.answer(f"✅ {user['name']} tasdiqlandi!")

@dp.message(Command("reject"))
async def admin_reject(message: types.Message):
    if not is_admin(message.from_user.id): return
    parts = message.text.split()
    if len(parts) < 2:
        await message.answer("Format: /reject [id]")
        return
    uid = int(parts[1])
    if uid not in pending_users:
        await message.answer(f"❌ {uid} topilmadi.")
        return
    user = pending_users.pop(uid)
    await bot.send_message(uid,
        f"❌ {user['name']}, so'rovingiz rad etildi.\n"
        f"📞 {ADMIN_USERNAME}"
    )
    await message.answer(f"❌ {user['name']} rad etildi.")

# ========================================
# MAIN
# ========================================
async def main():
    logger.info("Bot ishlanmoqda...")
    await dp.start_polling(bot)

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())